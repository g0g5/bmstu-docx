from __future__ import annotations

from pathlib import Path

from docx import Document

from bmstu_docx.docx_renderer import DocxRenderer
from bmstu_docx.models import HeadingBlock, InlineSpan, ParagraphBlock


def test_renderer_applies_page_setup_and_writes_docx(tmp_path: Path) -> None:
    blocks = [
        HeadingBlock(level=1, spans=[InlineSpan("ВВЕДЕНИЕ")], structural=True),
        ParagraphBlock(spans=[InlineSpan("Основной текст.")]),
    ]
    output_file = tmp_path / "output.docx"

    renderer = DocxRenderer()
    renderer.save(blocks, output_file)

    assert output_file.exists()

    document = Document(str(output_file))
    section = document.sections[0]
    assert round(section.left_margin.cm, 2) == 3.0
    assert round(section.right_margin.cm, 2) == 1.0
    assert document.paragraphs[0].text == "ВВЕДЕНИЕ"
    assert document.paragraphs[1].text == "Основной текст."
