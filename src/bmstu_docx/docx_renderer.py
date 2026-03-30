from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from .config import DEFAULT_CONFIG, DocumentConfig
from .docx_styles import (
    BLOCKQUOTE_STYLE,
    BODY_STYLE,
    CAPTION_STYLE,
    CODE_STYLE,
    FIGURE_TEXT_STYLE,
    HEADING1_STYLE,
    HEADING2_STYLE,
    HEADING3_STYLE,
    STRUCTURAL_HEADING_STYLE,
    TABLE_TEXT_STYLE,
    apply_document_styles,
    apply_run_font,
)
from .models import (
    Block,
    CodeBlock,
    HeadingBlock,
    ImageBlock,
    InlineSpan,
    ListBlock,
    ListItem,
    ParagraphBlock,
    TableBlock,
)


class DocxRenderer:
    def __init__(self, config: DocumentConfig = DEFAULT_CONFIG) -> None:
        self.config = config

    def render(self, blocks: list[Block]) -> Document:
        document = Document()
        apply_document_styles(document, self.config)
        for block in blocks:
            self._render_block(document, block, list_level=0)
        return document

    def save(self, blocks: list[Block], output_path: str | Path) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        document = self.render(blocks)
        document.save(str(output))
        return output

    def _render_block(self, document: Document, block: Block, list_level: int) -> None:
        if isinstance(block, ParagraphBlock):
            self._render_paragraph(document, block, list_level)
        elif isinstance(block, HeadingBlock):
            self._render_heading(document, block)
        elif isinstance(block, CodeBlock):
            self._render_code_block(document, block)
        elif isinstance(block, ListBlock):
            self._render_list(document, block, list_level)
        elif isinstance(block, TableBlock):
            self._render_table(document, block)
        elif isinstance(block, ImageBlock):
            self._render_image(document, block)

    def _render_paragraph(
        self, document: Document, block: ParagraphBlock, list_level: int = 0
    ) -> None:
        style_name = BODY_STYLE
        if block.kind == "blockquote":
            style_name = BLOCKQUOTE_STYLE
        elif block.kind == "figure_explanatory":
            style_name = FIGURE_TEXT_STYLE

        paragraph = document.add_paragraph(style=style_name)
        if list_level:
            paragraph.paragraph_format.left_indent = Cm(0.63 * list_level)
        self._render_spans(paragraph, block.spans)

    def _render_heading(self, document: Document, block: HeadingBlock) -> None:
        style_name = {
            1: HEADING1_STYLE,
            2: HEADING2_STYLE,
            3: HEADING3_STYLE,
        }.get(block.level, HEADING3_STYLE)
        if block.structural:
            style_name = STRUCTURAL_HEADING_STYLE
        paragraph = document.add_paragraph(style=style_name)
        self._render_spans(paragraph, block.spans, default_bold=True)

    def _render_code_block(self, document: Document, block: CodeBlock) -> None:
        lines = block.code.rstrip("\n").splitlines() or [""]
        for line in lines:
            paragraph = document.add_paragraph(style=CODE_STYLE)
            run = paragraph.add_run(line)
            apply_run_font(
                run, self.config.code_font_name, self.config.code_font_size_pt
            )

    def _render_list(
        self, document: Document, block: ListBlock, list_level: int
    ) -> None:
        style_name = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            self._render_list_item(document, item, style_name, list_level)

    def _render_list_item(
        self, document: Document, item: ListItem, style_name: str, list_level: int
    ) -> None:
        first_content_rendered = False
        for child in item.blocks:
            if isinstance(child, ParagraphBlock) and not first_content_rendered:
                paragraph = document.add_paragraph(style=style_name)
                paragraph.paragraph_format.left_indent = Cm(0.63 * list_level)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                self._render_spans(paragraph, child.spans)
                first_content_rendered = True
            elif isinstance(child, ListBlock):
                self._render_list(document, child, list_level + 1)
            else:
                self._render_block(document, child, list_level + 1)

        if not first_content_rendered:
            document.add_paragraph(style=style_name)

    def _render_table(self, document: Document, block: TableBlock) -> None:
        if block.caption is not None:
            caption = document.add_paragraph(style=CAPTION_STYLE)
            self._render_spans(caption, block.caption.spans)

        if not block.rows:
            return

        max_cols = max(len(row.cells) for row in block.rows)
        table = document.add_table(rows=len(block.rows), cols=max_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_index, row in enumerate(block.rows):
            for col_index, cell in enumerate(row.cells):
                table_cell = table.cell(row_index, col_index)
                paragraph = table_cell.paragraphs[0]
                paragraph.style = TABLE_TEXT_STYLE
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._render_spans(
                    paragraph,
                    cell.spans,
                    font_name=self.config.body_font_name,
                    font_size_pt=self.config.table_font_size_pt,
                )
                if row.header:
                    for run in paragraph.runs:
                        run.bold = True

    def _render_image(self, document: Document, block: ImageBlock) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_path = Path(block.path)
        if image_path.exists():
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(self.config.image_max_width_cm))
        else:
            run = paragraph.add_run(f"[Missing image: {image_path.name}]")
            apply_run_font(
                run, self.config.body_font_name, self.config.body_font_size_pt
            )

        if block.explanatory is not None:
            explanatory = ParagraphBlock(
                spans=block.explanatory.spans, kind="figure_explanatory"
            )
            self._render_paragraph(document, explanatory)

        if block.caption is not None:
            caption = document.add_paragraph(style=CAPTION_STYLE)
            self._render_spans(caption, block.caption.spans)

    def _render_spans(
        self,
        paragraph,
        spans: list[InlineSpan],
        default_bold: bool = False,
        font_name: str | None = None,
        font_size_pt: int | None = None,
    ) -> None:
        regular_font = font_name or self.config.body_font_name
        regular_size = font_size_pt or self.config.body_font_size_pt

        for span in spans:
            chunks = span.text.split("\n")
            for idx, chunk in enumerate(chunks):
                if idx > 0:
                    paragraph.add_run().add_break()
                if not chunk and len(chunks) > 1:
                    continue
                run = paragraph.add_run(chunk)
                if span.code:
                    apply_run_font(
                        run, self.config.code_font_name, self.config.code_font_size_pt
                    )
                else:
                    apply_run_font(run, regular_font, regular_size)
                run.bold = default_bold or span.bold
                run.italic = span.italic
